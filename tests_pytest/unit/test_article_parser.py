import tempfile
import os
from pathlib import Path
from core.article_parser import read_txt, read_pdf, read_docx, extract_article_text
import docx


def test_read_txt_with_utf8_encoding():
    content = "Ceci est un test avec accents éèê."
    content_bytes = content.encode("utf-8")

    # 1. Créer un fichier temporaire avec encodage connu
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as tmp:
        tmp.write(content_bytes)
        temp_path = Path(tmp.name)

    # 2. Lire le contenu avec la fonction testée
    result = read_txt(temp_path)

    # 3. Nettoyage
    os.remove(temp_path)

    # 4. Vérification
    assert isinstance(result, str)
    assert result == content


def generate_sample_pdf(text: str, filepath: Path):
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(str(filepath))
    c.drawString(100, 750, text)
    c.save()


def test_read_pdf_extracts_text():
    content = "Hello PDF world!"

    # 1. Créer un PDF temporaire
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        temp_path = Path(tmp.name)
        generate_sample_pdf(content, temp_path)

    # 2. Lire le contenu via la fonction
    result = read_pdf(temp_path)

    # 3. Nettoyage
    os.remove(temp_path)

    # 4. Vérifications
    assert isinstance(result, str)
    assert content in result


def test_read_docx_extracts_all_paragraphs():
    content = ["Bonjour", "Voici un test", "Fin"]

    # 1. Créer un fichier Word temporaire avec plusieurs paragraphes
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = Path(tmp.name)
        document = docx.Document()
        for line in content:
            document.add_paragraph(line)
        document.save(temp_path)

    # 2. Lire le contenu via la fonction
    result = read_docx(temp_path)

    # 3. Nettoyage
    os.remove(temp_path)

    # 4. Vérifications
    assert isinstance(result, str)
    assert result == "\n".join(content)


def test_extract_article_text_with_txt():
    content = "Ceci est un .txt"
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
        tmp.write(content)
        temp_path = Path(tmp.name)
    result = extract_article_text(temp_path)
    os.remove(temp_path)
    assert result == content


def test_extract_article_text_with_docx():
    content = ["Ligne 1", "Ligne 2"]
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = Path(tmp.name)
        document = docx.Document()
        for line in content:
            document.add_paragraph(line)
        document.save(temp_path)
    result = extract_article_text(temp_path)
    os.remove(temp_path)
    assert result == "\n".join(content)


def test_extract_article_text_raises_on_unknown_extension():
    with tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode="w", encoding="utf-8") as tmp:
        tmp.write("# Titre Markdown")
        temp_path = Path(tmp.name)

    try:
        extract_article_text(temp_path)
        assert False, "Une exception aurait dû être levée"
    except ValueError as e:
        assert "Format non pris en charge" in str(e)
    finally:
        os.remove(temp_path)


def test_extract_article_text_with_pdf():
    content = "Contenu PDF de test"
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        temp_path = Path(tmp.name)
        generate_sample_pdf(content, temp_path)

    result = extract_article_text(temp_path)
    os.remove(temp_path)

    assert isinstance(result, str)
    assert content in result
